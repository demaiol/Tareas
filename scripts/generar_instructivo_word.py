from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

output_path = "/Users/leonardodemaio/Library/Mobile Documents/com~apple~CloudDocs/Codex Local/Tareas/Instructivo_Uso_Aplicacion_Tareas.docx"


doc = Document()

# Titulo
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Instructivo de Uso - Aplicación Tareas")
run.bold = True
run.font.size = Pt(18)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.add_run(f"Versión: {datetime.now().strftime('%d-%m-%Y')}")


doc.add_paragraph()
intro = doc.add_paragraph()
intro.add_run("Objetivo: ").bold = True
intro.add_run(
    "Explicar de forma simple cómo utilizar los módulos Administrador y Reportes "
    "para gestionar requerimientos recibidos por correo."
)

# URLs
doc.add_heading("1. Acceso a la aplicación", level=2)
doc.add_paragraph("- Módulo Reportes (solo lectura): https://tareas-reportes.streamlit.app")
doc.add_paragraph("- Módulo Administrador (gestión): [COMPLETAR URL DEL ADMINISTRADOR]")
doc.add_paragraph("- Credenciales Reportes: usuario 'gestion' y contraseña 'gestion123$'.")
doc.add_paragraph("- Credenciales Administrador: usuario 'Administrador' y contraseña 'DEMO123$'.")

img1 = doc.add_paragraph("[PEGAR IMAGEN 1: Pantalla de login del módulo Reportes]")
img1.runs[0].italic = True
img2 = doc.add_paragraph("[PEGAR IMAGEN 2: Pantalla de login del módulo Administrador]")
img2.runs[0].italic = True

# Administrador
doc.add_heading("2. Uso del módulo Administrador", level=2)
doc.add_paragraph("1. Ingresar al módulo Administrador y autenticarse.")
doc.add_paragraph("2. Presionar el botón 'Sincronizar correos no leídos' para importar nuevos requerimientos.")
doc.add_paragraph("3. Revisar la tabla de requerimientos y seleccionar una fila (REQ) para abrir el detalle.")
doc.add_paragraph("4. Completar respuesta, actualizar estado y guardar.")
doc.add_paragraph("5. Si el estado pasa a 'Resuelto', se enviará automáticamente un correo de cierre al solicitante.")

img3 = doc.add_paragraph("[PEGAR IMAGEN 3: Pantalla principal de Administrador con botón 'Sincronizar correos no leídos']")
img3.runs[0].italic = True
img4 = doc.add_paragraph("[PEGAR IMAGEN 4: Tabla de REQ y selección de una fila]")
img4.runs[0].italic = True
img5 = doc.add_paragraph("[PEGAR IMAGEN 5: Formulario de actualización de requerimiento]")
img5.runs[0].italic = True

# Reportes
doc.add_heading("3. Uso del módulo Reportes", level=2)
doc.add_paragraph("1. Ingresar a https://tareas-reportes.streamlit.app con credenciales de Reportes.")
doc.add_paragraph("2. Revisar indicadores generales (Total, Pendientes, Resueltos, % de resolución).")
doc.add_paragraph("3. Analizar gráficos de estado y tendencias de temas pendientes/resueltos.")
doc.add_paragraph("4. Seleccionar un REQ en la tabla para ver el detalle y la resolución registrada (solo lectura).")

img6 = doc.add_paragraph("[PEGAR IMAGEN 6: Dashboard del módulo Reportes con KPIs y gráficos]")
img6.runs[0].italic = True
img7 = doc.add_paragraph("[PEGAR IMAGEN 7: Detalle del REQ y sección de resolución en Reportes]")
img7.runs[0].italic = True

# Buenas practicas
doc.add_heading("4. Buenas prácticas operativas", level=2)
doc.add_paragraph("- Sincronizar correos al inicio de la jornada y al cierre del día.")
doc.add_paragraph("- Registrar respuestas claras y completas al resolver un requerimiento.")
doc.add_paragraph("- Verificar en Reportes los pendientes para priorizar gestión.")
doc.add_paragraph("- Mantener credenciales de acceso en reserva y cambiarlas periódicamente.")

# Contacto / soporte
doc.add_heading("5. Soporte", level=2)
doc.add_paragraph("Ante fallas de acceso, sincronización o envío de correos, contactar al administrador técnico de la plataforma.")


doc.save(output_path)
print(output_path)
