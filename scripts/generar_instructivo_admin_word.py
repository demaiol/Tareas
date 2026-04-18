from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

output_path = "/Users/leonardodemaio/Library/Mobile Documents/com~apple~CloudDocs/Codex Local/Tareas/Instructivo_Modulo_Administrador.docx"

doc = Document()

# Titulo
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Instructivo de Uso - Módulo Administrador")
run.bold = True
run.font.size = Pt(18)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.add_run(f"Versión: {datetime.now().strftime('%d-%m-%Y')}")

doc.add_paragraph()
intro = doc.add_paragraph()
intro.add_run("Objetivo: ").bold = True
intro.add_run(
    "Explicar de forma simple cómo operar el módulo Administrador "
    "para gestionar requerimientos recibidos por correo."
)

# Acceso
doc.add_heading("1. Acceso al módulo Administrador", level=2)
doc.add_paragraph("- URL del módulo Administrador: [COMPLETAR URL DEL ADMINISTRADOR]")
doc.add_paragraph("- Credenciales: usuario 'Administrador' y contraseña 'DEMO123$'.")
img1 = doc.add_paragraph("[PEGAR IMAGEN 1: Pantalla de login del módulo Administrador]")
img1.runs[0].italic = True

# Operacion diaria
doc.add_heading("2. Operación diaria", level=2)
doc.add_paragraph("1. Ingresar al módulo Administrador y autenticarse.")
doc.add_paragraph("2. Presionar el botón 'Sincronizar correos no leídos' para importar solicitudes nuevas.")
doc.add_paragraph("3. Revisar el reporte operativo y seleccionar una fila (REQ) para abrir el detalle.")
doc.add_paragraph("4. Completar respuesta, actualizar estado y guardar.")
doc.add_paragraph("5. Si cambia a 'Resuelto', se enviará automáticamente un correo de cierre al solicitante.")

img2 = doc.add_paragraph("[PEGAR IMAGEN 2: Pantalla principal con botón 'Sincronizar correos no leídos']")
img2.runs[0].italic = True
img3 = doc.add_paragraph("[PEGAR IMAGEN 3: Tabla de REQ y selección de una fila]")
img3.runs[0].italic = True
img4 = doc.add_paragraph("[PEGAR IMAGEN 4: Formulario de actualización del requerimiento]")
img4.runs[0].italic = True

# Correos automaticos
doc.add_heading("3. Correos automáticos", level=2)
doc.add_paragraph("- Al crear un REQ nuevo, se envía acuse al remitente con número de requerimiento.")
doc.add_paragraph("- Al cerrar un REQ (estado Resuelto), se envía correo de cierre con detalle de resolución.")
img5 = doc.add_paragraph("[PEGAR IMAGEN 5: Mensajes de confirmación de sincronización y envío de correo]")
img5.runs[0].italic = True

# Buenas practicas
doc.add_heading("4. Buenas prácticas", level=2)
doc.add_paragraph("- Ejecutar sincronización al inicio y al cierre de la jornada.")
doc.add_paragraph("- Mantener respuestas claras y completas para cada cierre.")
doc.add_paragraph("- Usar filtros de estado para priorizar pendientes.")
doc.add_paragraph("- Resguardar credenciales de acceso y cambiarlas periódicamente.")

# Soporte
doc.add_heading("5. Soporte", level=2)
doc.add_paragraph("Ante fallas de acceso, sincronización o envío de correos, contactar al administrador técnico.")

doc.save(output_path)
print(output_path)
